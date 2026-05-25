"""
Build Final Project Report (11).docx by combining:
  - Final Project Report (10).docx — base
  - Final Project Writing .docx    — full prose + 11 figures (other-github writing)
  - finalized_outputs/*            — our model figures (Phase A, Phase C, K-sweep)
                                     + cross-snapshot summary

Layout: paste writing-doc prose and images into the empty Results/Analyses
subsections; layer our model figures with 1-3 sentence captions; relocate the
Cross-Method Comparison block from its current standalone slot into
Analyses > Tokens/compute.
"""
from __future__ import annotations
import os, shutil, copy
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
from docx.text.paragraph import Paragraph


REPO       = Path(r'C:\Users\WCCO\Documents\GitHub\FTG Final')
SRC_REPORT = Path(r'D:\archive\Final Project Report (10).docx')
SRC_WRITE  = Path(r'D:\archive\Final Project Writing .docx')
DST_REPORT = Path(r'D:\archive\Final Project Report (16).docx')
ASSETS     = REPO / 'finalized_outputs' / 'report_v11_assets'
WRITE_IMGS = ASSETS / 'writing_doc_images'


# ── verbatim prose blocks pulled from Final Project Writing .docx ──────────
# Keys = writing-doc paragraph indices used in extraction earlier.
WRITING_PROSE = {
    # image2 (LMP vs toll daily mean, 3 months) — P8-11
    'lmp_vs_toll_daily': [
        "Over the three months shown, daily-mean LMPs at HB_HOUSTON ranged from roughly $20 to $60/MWh, while the corresponding toll cost sat between $50 and $80/MWh and trended upward with rising gas prices. On most days, the average grid price fell well below the toll line under both gas conventions, meaning that a straightforward grid purchase would have been the lower-cost option the majority of the time. This held true regardless of whether we modeled fuel costs using HSC (top panel) or Henry Hub (bottom panel), though the Henry Hub-based toll was generally a few dollars higher, widening the gap further.",
        "However, the shaded band, which extends from the daily mean up to the 95th-percentile of hourly LMPs, frequently pierces above the toll cost line, particularly during what are likely evening peak hours. These spikes, sometimes reaching $80–$150/MWh, represent intervals where self-generating under the toll would have been cheaper than buying from the grid. This pattern was especially pronounced in early and mid-November, when several days showed the p95 band extending to well over $100/MWh, far exceeding the toll cost on those same days.",
        "This is where optionality becomes valuable. Rather than committing to one strategy full-time, having the flexibility to switch from grid purchases to tolled generation during peak hours allows us to capture the best price in both regimes. The upper edge of the band essentially signals when the toll option moves \"into the money,\" and the frequency with which it does so suggests that this optionality carries meaningful economic value even though the grid wins on a daily-average basis. In practical terms, the ideal approach would be to buy from the grid during off-peak and shoulder hours when LMPs are low, and dispatch the tolling asset during the handful of peak hours when prices spike above the cost of self-generation.",
        "It is also worth noting that as tolling costs trended upward over the period driven by rising gas prices heading into winter, the threshold for the toll to be competitive moved higher as well. This means the value of optionality is not static; it depends on the prevailing fuel environment and the volatility of real-time power prices. In a higher-gas-price environment, the toll becomes harder to justify on average, but if grid price spikes grow more severe alongside tighter winter conditions, the peak-hour payoff from tolling could increase in tandem.",
    ],
    # image8 (hourly LMP vs toll + spark spread win %) — P16-18
    'lmp_vs_toll_hourly': [
        "To better understand when tolling actually makes sense, we broke down the LMP and toll cost comparison by hour of day over the past three months. The left panel plots mean hourly LMPs against the mean toll cost using HSC gas pricing. What stands out immediately is that the toll cost is basically flat across the day around $58-60/MWh every hour because it is tied to daily gas prices which don’t fluctuate diurnally nearly as much as real-time electricity prices. Electricity prices, on the other hand, follow a clear daily cycle. They sit in the $20–35/MWh range for most of the day and then spike sharply in the late afternoon, hitting roughly $65/MWh around hours 17 and 18. That evening peak is the only window where mean LMPs consistently approach or exceed the toll cost.",
        "The right panel above puts a finer point on this by showing the percentage of hours where the toll actually beat the grid (i.e., a positive spark spread). For most hours of the day, tolling wins less than 10% of the time so the grid is almost always cheaper. But starting around hour 16, the numbers climb quickly, reaching about 55–60% at peak under HSC gas and slightly less under Henry Hub. By hours 20-22, the share tapers off but still hovers around 25–35% before falling back overnight.",
        "For us, the takeaway is straightforward. The toll is not something we would want to run around the clock. The grid is just too cheap during off-peak hours. But during that five- to six-hour evening window, tolling becomes a genuinely useful hedge against price spikes. We think the right way to use it is as an option we exercise selectively during peak hours, which also has implications for when we schedule training loads and how we structure our overall power procurement.",
    ],
    # image6 (heatmaps for both locations) — P22-24
    'lmp_heatmaps': [
        "To inform our power procurement strategy, we plotted heatmaps of mean hourly LMPs across all months of 2025 for both data center locations. The axes show month on the horizontal and hour of day on the vertical, with color intensity representing average price levels in $/MWh.",
        "The most obvious pattern in both locations is the concentration of high prices during evening hours, roughly between hours 16 and 22, with the most extreme values reaching $80-100/MWh and clustered in the summer months (June through August) around hours 18-20. This aligns with what we would expect- peak electricity demand from cooling loads and tightening supply during hot Texas afternoons and evenings. Outside of that window, prices drop substantially.",
        "From a practical standpoint, this chart helped us think about when and where to schedule training loads. The deep blue zones represent the cheapest windows to run power-intensive tasks from the grid - between hours 8 and 16. This dip in prices during the day is likely driven by the penetration of solar resources in Texas. The summer evening peaks, on the other hand, are the hours where we would want to either curtail grid exposure or lean on the tolling or BESS options. The wide spread between peak and off-peak prices also suggests that battery storage arbitrage could be particularly attractive, which we explore further.",
    ],
    # image1 (IMHR distribution at HB_HOUSTON) — P30-32
    'imhr_distribution': [
        "To evaluate whether our tolling asset would be economically dispatched in a given hour, we computed the implied market heat rate (IMHR) at HB_HOUSTON for each hour over the last three months of 2025. The IMHR is calculated as the hourly LMP divided by the all-in fuel cost, where the fuel cost equals the relevant daily gas price plus a $3/MMBtu adder for variable operations and maintenance (VOM). A higher IMHR means the electricity market is pricing power more generously relative to gas, and when the IMHR exceeds the physical heat rate of a given generator, that unit would be \"in the money\" (ITM) to dispatch. We show two versions: the left panel uses Houston Ship Channel (HSC) gas pricing, which better reflects local delivered costs in Texas, while the right panel uses Henry Hub (HH) pricing per the project specification.",
        "Both distributions are clearly right-skewed, with the bulk of observations clustered between roughly 2 and 8 MMBtu/MWh. Under the Henry Hub convention, the median IMHR comes in at 4.56 MMBtu/MWh, and under HSC it is slightly higher at 4.80. We overlaid three reference lines corresponding to different generation technologies - the CCGT at 6.8 MMBtu/MWh, the Antelope Station peaker at 9.2, and a simple-cycle peaker at 9.5, which matches the heat rate assumed in our tolling contract. Under the HH specification, about 20.0% of hours exceed the CCGT threshold and 7.7% exceed the 9.5 peaker line. The HSC panel shows slightly higher ITM percentages (23.7% for CCGT and 8.8% for the peaker), which reflects the difference in gas pricing conventions between the two panels.",
        "What matters most for our analysis is the long right tail. Both distributions extend past 10 MMBtu/MWh, and while these extreme hours are relatively rare, they represent scarcity conditions where grid prices are higher than the generation cost. These are the hours that drive the option value of the tolling contract. Even though the median IMHR sits well below the 9.5 peaker threshold, meaning the toll would be out of the money most of the time, the handful of hours in the tail generate outsized positive spark spreads that can offset many hours of negative or zero payoff. This is consistent with what we observed in the diurnal analysis - the toll is not a baseload play but an option on price spikes, and this distribution helps us visualize exactly how often and how deeply that option pays off.",
    ],
    # image10 (intermittency: Houston quantile + West negative %) — P36-39
    'intermittency': [
        "The charts above capture what we see as the structural basis for the toll option's value: the role of renewable intermittency in creating asymmetric price behavior across both locations.",
        "The left panel shows the distribution of hourly LMPs at HB_HOUSTON broken out by hour of day, with quantile bands computed across all dates in the sample. The median and mean lines follow the familiar diurnal pattern we discussed earlier, sitting in the $25-35/MWh range for most of the day and rising to about $50-65/MWh during evening hours. What is more telling, though, is the shape of the tails. During daytime and overnight hours, the quantile bands are relatively tight, meaning prices are fairly predictable. But starting around hour 16 and extending through hour 22, the 95th-to-99th percentile band blows out dramatically, reaching upward of $150–200/MWh. This shows that a normal evening hour is not particularly remarkable, but a \"bad\" evening hour (from the buyer's perspective) can see prices spike to five or even ten times the typical level. That skew is exactly what gives the tolling option its value. Most hours will expire worthless, but the rare extreme hours generate outsized payoffs that more than compensate.",
        "The right panel shifts to HB_WEST and shows the percentage of hours with negative LMPs by month, which we interpret as a wind oversupply signature. March stands out with roughly 15% of hours going negative, followed by April at about 9%. The spring concentration makes sense given that during those months weather and resulting residential demand are still relatively mild. Other months generally fall in the 1–5% range, though there is a modest uptick in the fall around October and November. Negative prices are a direct consequence of renewable intermittency. When wind output floods the grid faster than demand, prices get pushed below zero.",
        "Taken together, these two panels illustrate both sides of the intermittency coin. In Houston, renewable variability contributes to extreme upward price spikes during evening hours when solar output drops off and demand remains high. In West Texas, the same underlying intermittency manifests as negative prices during periods of wind oversupply. For our analysis, the left panel is the more actionable one because it shows the conditions under which our toll option pays off. But the right panel is also relevant for thinking about when to schedule training loads at the West location, since negative price hours represent periods where we are effectively being paid to consume power.",
    ],
    # image11 (Monte-Carlo 5000 paths, 4 state vars) — P46-49
    'mc_validation': [
        "To project operating economics over the June through November 2026 period, we ran a Monte Carlo simulation with 5,000 paths for each of our four key state variables: hourly power prices at HB_HOUSTON and HB_WEST, and daily natural gas prices at both HSC and Henry Hub. The chart above plots 50 representative paths for each variable along with the 5th-to-95th percentile fan across all 5,000 simulations. The dashed red line in each panel marks the historical mean for reference.",
        "Starting with the two power price panels, the simulated paths for both Houston and West show median trajectories hovering around $25–35/MWh, which is broadly consistent with the historical means of $35.0 and $34.1 respectively. What stands out in both panels, though, is the frequency and magnitude of upward spikes. Individual paths regularly shoot above $100/MWh, and the upper edge of the confidence band reaches roughly $250/MWh in Houston and nearly $400/MWh in West. These are the types of scarcity hours we identified in the IMHR and diurnal analyses, and the simulation confirms that they are a persistent feature of the price landscape rather than one-off anomalies. The West panel also shows a slightly wider band overall, which is consistent with the greater price volatility we observed at that location due to wind intermittency.",
        "The bottom two panels show the gas price simulations, and the behavior here is quite different. Rather than the spiky, mean-reverting pattern we see in power prices, the gas paths exhibit a gradually widening fan over time. Both HSC and Henry Hub start the period tightly clustered near their historical means ($3.1/MMBtu and $3.5/MMBtu, respectively), but by October and November the 5th-to-95th percentile band expands considerably, with upper-end paths reaching $9–10/MMBtu or higher. This widening reflects the compounding uncertainty in gas prices as we move further into the future, and the upward skew toward the fall months captures the seasonal risk of rising heating demand heading into winter. The practical implication is that tolling costs become less predictable later in the operating period, which adds another layer of uncertainty to the dispatch decision.",
        "One thing we want to flag is the interaction between these variables. The toll option pays off when power prices spike while gas prices remain moderate, since that is when the spark spread is widest. But if gas prices also happen to be elevated during a power spike (as might occur during a winter weather event, for example), the toll cost rises in tandem and the payoff shrinks. The simulation captures this correlation structure, and it is something we account for in our dispatch optimization and option valuation results that follow.",
    ],
}

# Orphan images from the writing doc with no prose written yet
ORPHAN_IMAGES = {
    'valuation_1':       WRITE_IMGS / 'image9.png',
    'valuation_2':       WRITE_IMGS / 'image3.png',
    'location_compare_1': WRITE_IMGS / 'image4.png',
    'location_compare_2': WRITE_IMGS / 'image5.png',
    'risk_analysis':     WRITE_IMGS / 'image7.png',
}


# ── helper: insert paragraph immediately after `prev_para` ─────────────────
def insert_para_after(prev_para, doc, *, text='', style='Normal', bold=False,
                       italic=False, size_pt=None, alignment=None):
    """Insert a new <w:p> after `prev_para`; return wrapped Paragraph."""
    new_p_xml = OxmlElement('w:p')
    prev_para._element.addnext(new_p_xml)
    new_para = Paragraph(new_p_xml, prev_para._parent)
    if style in (s.name for s in doc.styles):
        new_para.style = doc.styles[style]
    if text:
        run = new_para.add_run(text)
        if bold:   run.bold = True
        if italic: run.italic = True
        if size_pt: run.font.size = Pt(size_pt)
    if alignment is not None:
        new_para.alignment = alignment
    return new_para


def insert_image_para_after(prev_para, doc, image_path, width_in=6.0,
                             caption=None):
    """Insert a centred-image paragraph (and optionally a caption) after prev."""
    img_para = insert_para_after(prev_para, doc, style='Normal',
                                  alignment=WD_ALIGN_PARAGRAPH.CENTER)
    img_para.add_run().add_picture(str(image_path), width=Inches(width_in))
    last = img_para
    if caption:
        cap = insert_para_after(last, doc, text=caption, style='Normal',
                                 italic=True, size_pt=10,
                                 alignment=WD_ALIGN_PARAGRAPH.CENTER)
        last = cap
    return last


def _apply_table_borders(table, sz=4, color='888888'):
    """Apply single-line borders to all cells of a python-docx Table."""
    tblPr = table._element.find(qn('w:tblPr'))
    if tblPr is None:
        tblPr = OxmlElement('w:tblPr')
        table._element.insert(0, tblPr)
    borders = OxmlElement('w:tblBorders')
    for side in ('top','left','bottom','right','insideH','insideV'):
        b = OxmlElement(f'w:{side}')
        b.set(qn('w:val'), 'single')
        b.set(qn('w:sz'),  str(sz))
        b.set(qn('w:space'), '0')
        b.set(qn('w:color'), color)
        borders.append(b)
    tblPr.append(borders)


def insert_table_after(prev_para, doc, rows_data, header=True, widths_in=None):
    """Insert a table after `prev_para`. rows_data is list-of-lists (str cells).
    `header=True` makes the first row bold."""
    # Create table at end of doc, then move it
    n_rows, n_cols = len(rows_data), len(rows_data[0]) if rows_data else 0
    tbl = doc.add_table(rows=n_rows, cols=n_cols)
    # Use TableNormal (always present); apply borders via direct XML for grid.
    tbl.style = doc.styles['TableNormal']
    _apply_table_borders(tbl)
    for r_idx, row in enumerate(rows_data):
        for c_idx, val in enumerate(row):
            cell = tbl.cell(r_idx, c_idx)
            cell.text = str(val)
            if header and r_idx == 0:
                for p in cell.paragraphs:
                    for r in p.runs:
                        r.bold = True
    if widths_in:
        for r in tbl.rows:
            for i, c in enumerate(r.cells):
                if i < len(widths_in):
                    c.width = Inches(widths_in[i])
    # Move table element after prev_para
    prev_para._element.addnext(tbl._element)
    # Return last element handle (still the table)
    return tbl


def append_after_table(table_elem, parent_proxy, doc, *, text='', style='Normal',
                       italic=False, alignment=None, size_pt=None):
    """Insert a paragraph immediately after a table element.
    parent_proxy must be a known-good _parent (e.g. an existing paragraph's _parent)."""
    new_p_xml = OxmlElement('w:p')
    table_elem._element.addnext(new_p_xml)
    new_para = Paragraph(new_p_xml, parent_proxy)
    if style in (s.name for s in doc.styles):
        new_para.style = doc.styles[style]
    if text:
        run = new_para.add_run(text)
        if italic: run.italic = True
        if size_pt: run.font.size = Pt(size_pt)
    if alignment is not None:
        new_para.alignment = alignment
    return new_para


# ── main ──────────────────────────────────────────────────────────────────
def main():
    print(f'Opening base report: {SRC_REPORT}')
    doc = Document(str(SRC_REPORT))

    # 1) Find subsection heading paragraphs (single pass — cache live Paragraph objs
    #    and track Cross-Method body paragraphs in the same walk)
    targets = {}
    cross_method_paras = []
    in_cross_method = False
    ctx = None
    for p in doc.paragraphs:
        txt = p.text.strip()
        if txt == 'Conclusion':
            in_cross_method = False
            ctx = 'Conclusion'
            targets['conclusion'] = p
        elif txt.startswith('Cross-Method Comparison'):
            ctx = 'CrossMethod'
            targets['cross_method_heading'] = p
            in_cross_method = True
            continue
        elif txt == 'Results':
            in_cross_method = False
            ctx = 'Results'
        elif txt == 'Analyses and Sensitivities':
            in_cross_method = False
            ctx = 'Analyses'
        if in_cross_method and txt:
            cross_method_paras.append(p)
        if txt in ['Power pricing', 'Tokens/compute',
                    'BESS/Nat-gas Tolling Optionality'] and ctx in ('Results', 'Analyses'):
            targets[(ctx, txt)] = p

    print(f'Found {len(cross_method_paras)} cross-method body paragraphs to relocate.')

    # 2) Insert content per subsection, in DOCUMENT ORDER so we don't trip ourselves.
    #    Use insert_para_after with chained anchors.

    # ───────── Final Policy callout — inserted right after "Results" heading ─────
    # Locate the Results heading (it's the paragraph just before Results > Power pricing)
    pp_anchor = targets[('Results', 'Power pricing')]
    # Walk back to find the "Results" heading paragraph
    body = pp_anchor._element.getparent()
    children = list(body)
    pp_idx = children.index(pp_anchor._element)
    results_heading_elem = None
    for j in range(pp_idx - 1, -1, -1):
        if children[j].tag == qn('w:p'):
            txt = ''.join(t.text or '' for t in children[j].iter(qn('w:t'))).strip()
            if txt == 'Results':
                results_heading_elem = children[j]
                break
    if results_heading_elem is not None:
        results_para = Paragraph(results_heading_elem, pp_anchor._parent)
        cb_h = insert_para_after(results_para, doc,
            text='Final Policy', style='Heading 3')
        # Opening paragraph — robust across all 16 drift × stress combinations.
        cb_b = insert_para_after(cb_h, doc, style='Normal',
            text=('Across all four committed drift scenarios (baseline, ai_structural, '
                  'mild_drift, ai_plus_brent) and all four stress overlays (no-stress, '
                  'mild, moderate, uri_full), the LP converges on a 90-day training cadence '
                  '(verified under Phase C optimal procurement in every snapshot). Procurement '
                  'choice, however, is stress-dependent. The four rungs of the stress ladder are:'))
        # Stress-ladder bullets — one per overlay.
        cb_b = insert_para_after(cb_b, doc, style='Normal',
            text=('  • No-stress (none).  LMP-only is optimal at every drift level. Mean '
                  'profit $36,377M (baseline) → $36,376M (ai_plus_brent). Toll wins by $2.18M '
                  'at K=$0 but the buyer\'s breakeven is only K* ≈ $3.6/kW-mo — far below the '
                  'seller\'s $8/kW-mo anchor — so optimal MW reservation collapses to 0 at the '
                  'seller rate. Toll value is real but consumed entirely by the lease.'))
        cb_b = insert_para_after(cb_b, doc, style='Normal',
            text=('  • Mild scarcity (72h spike, $200-$400/MWh, p=0.50).  LMP-only still '
                  'wins at the default K=$8. Mean profit $36,376M (baseline) → $36,374M '
                  '(ai_plus_brent), down ~$1.6M vs no-stress because the spike adds LMP cost '
                  'that the LP cannot fully avoid. The toll-vs-LMP-only gap closes from '
                  '−$2.62M to −$1.97M but does not flip.'))
        cb_b = insert_para_after(cb_b, doc, style='Normal',
            text=('  • Moderate scarcity (96h spike, $500-$1.5K/MWh + $20 HH gas, p=0.20).  '
                  'LMP-only still wins. Mean profit $36,374M (baseline) → $36,372M (ai_plus_brent), '
                  'a further ~$1.5M drop. Toll-vs-LMP-only gap closes another $0.7M to −$1.27M but '
                  'remains negative — at the seller\'s $8/kW-mo rate the lease still dominates the '
                  'dispatch savings even with $1.5K/MWh shoulder spikes.'))
        cb_b = insert_para_after(cb_b, doc, style='Normal',
            text=('  • Uri-style scarcity (100h spike, $5K-$9K/MWh + $250 HH gas, p=0.05).  '
                  'Optimal procurement flips to LMP + toll in every drift scenario. Mean profit '
                  '$36,366M (baseline) → $36,364M (ai_plus_brent); toll wins by +$2.27M to +$2.39M '
                  'at the default K=$8 (cadence-invariant — confirmed against 30d K-sweep). The '
                  'buyer\'s breakeven K* climbs to ≈ $11.8/kW-mo, well above the seller\'s '
                  '$8/kW-mo anchor — the toll becomes economically signable at market rate.'))
        # Closing paragraph — pulls the four rungs into one decision recommendation.
        cb_b = insert_para_after(cb_b, doc, style='Normal',
            text=('Bottom line. Drift sensitivity is small (~$1.7M across the four drifts at any '
                  'fixed stress); the stress overlay is the dominant axis (~$1.5M / step through '
                  'mild/moderate, ~$11M jump at uri_full). The toll is a tail-insurance instrument: '
                  'its value at the default $8/kW-mo seller rate only exceeds the lease cost under '
                  'a Uri-class winter-storm scenario. Recommendation: sign LMP-only under expected '
                  'conditions; revisit toll procurement if scarcity-risk priors shift toward Uri-style '
                  'tail events (e.g., updated ERCOT capacity adequacy projections, gas-supply stress).'))

    # ───────── Results > Power pricing ─────────
    anchor = targets[('Results', 'Power pricing')]
    anchor = insert_para_after(anchor, doc,
        text='ERCOT LMP heatmaps for HB_HOUSTON and HB_WEST',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        WRITE_IMGS / 'image6.png', width_in=6.0,
        caption='Figure: Mean hourly LMP heatmaps by (month × hour-of-day) for both data center locations, full-year 2025 ERCOT DAM data.')
    for para_text in WRITING_PROSE['lmp_heatmaps']:
        anchor = insert_para_after(anchor, doc, text=para_text, style='Normal')
    # Orphan location-comparison images
    anchor = insert_para_after(anchor, doc, text='Location comparison (additional)',
                                style='Heading 3')
    for key in ['location_compare_1', 'location_compare_2']:
        anchor = insert_image_para_after(anchor, doc, ORPHAN_IMAGES[key],
            width_in=6.0,
            caption='Figure: Location comparison diagnostic (analysis text pending).')

    # ───────── Results > Tokens/compute ─────────
    anchor = targets[('Results', 'Tokens/compute')]
    anchor = insert_para_after(anchor, doc,
        text='Phase A — Training cadence selection',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'phase_a_cadence_sweep.png', width_in=6.0,
        caption='Figure: Phase A mean profit ($M) by candidate training cadence across 50 MC paths under doc_blended scheme + 0.25 global revenue haircut, baseline drift. The 90-day cadence wins at $36,372M and is verified under Phase C optimal procurement.')
    anchor = insert_para_after(anchor, doc,
        text='Phase A scores each candidate cadence by mean profit across 50 Monte-Carlo price paths under the doc_blended token-multiplier scheme with the 0.25 global haircut applied. Profit grows monotonically with cadence because longer cadences spend less compute-MWh on (zero-revenue) training and each release window\'s compute requirement grows with release date as projected model sizes scale.',
        style='Normal')
    # NEW: Per-release training schedule at the winning 90d cadence
    anchor = insert_para_after(anchor, doc,
        text='Per-release training schedule at the winning 90d cadence',
        style='Heading 3')
    schedule_rows = [
        ['Release', 'Training-window start', 'Release date', 'Compute required (cMWh)', '% of 6-mo total', 'Multiplier'],
        ['R1 (initial)', '2026-06-01', '2026-06-11', '34,869', '27.7%', '0.328'],
        ['R2',           '2026-06-11', '2026-09-09', '36,740', '29.2%', '0.350'],
        ['R3',           '2026-09-09', '2026-12-08', '54,228', '43.1%', '0.374'],
    ]
    parent_proxy = anchor._parent
    sched_table = insert_table_after(anchor, doc, schedule_rows,
                                      widths_in=[0.9, 1.4, 1.2, 1.6, 1.1, 0.9])
    anchor = append_after_table(sched_table, parent_proxy, doc,
        text='Table: At the 90-day cadence, three model releases fit inside the Jun-Nov 2026 horizon. R1 is the initial release (100% compute → training, 10-day fast-finish), R2 and R3 are standard 60-90 day cadence releases. Per-release compute requirements grow with release date as projected frontier model size scales. The token-revenue multiplier column shows the composed doc_blended × 0.25 global haircut effect for each release window.',
        style='Normal', italic=True, size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # NEW: Diurnal training/inference allocation pattern
    anchor = insert_para_after(anchor, doc,
        text='Diurnal training and inference allocation pattern (LP output)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'train_inf_diurnal.png', width_in=6.0,
        caption='Figure: Hour-of-day mean allocation across both sites at the 90d cadence × LMP-only winning policy. Training is concentrated in low-LMP overnight and midday hours; inference dominates the evening peak window when revenue per grid-MWh ($166K) far exceeds even the highest spikes in marginal power cost.')

    # ───────── Results > BESS/Nat-gas Tolling Optionality ─────────
    anchor = targets[('Results', 'BESS/Nat-gas Tolling Optionality')]
    anchor = insert_para_after(anchor, doc,
        text='Phase C — Procurement comparison at the locked 90d cadence',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'phase_c_procurement_bars_twopanel.png', width_in=6.8,
        caption='Figure: 8-combination procurement comparison at 90d cadence, two panels — left = no-stress baseline drift (LMP-only wins at $36,377M); right = Uri-stress overlay (LMP + toll wins at $36,366M, with +$2.27M margin over LMP-only). Bars show Δ profit vs the panel-specific winner; $8/kW-mo toll capacity payment and $3M/site BESS lease are deducted from all toll/BESS scenarios. The winner-flip from LMP-only → LMP+toll under Uri-stress is the central optionality finding.')
    anchor = insert_para_after(anchor, doc,
        text='Under the baseline drift with no Uri-style stress overlay, the LMP-only procurement choice is the Phase C winner. The flatness across the eight scenarios under no-stress is the headline result: at current market lease rates ($8/kW-mo toll, $3M/site BESS), the gross dispatch value the LP captures from optionality is not large enough to clear the fixed leases. Under the Uri-stress overlay, the toll\'s tail-insurance value rises sharply and LMP + toll becomes the dominant policy across every drift scenario (see the cross-snapshot drift summary in the Analyses section).',
        style='Normal')
    # NEW: Physical vs Virtual BESS comparison (addresses Deliverable 2 BESS-at-each-location)
    anchor = insert_para_after(anchor, doc,
        text='Physical vs Virtual BESS comparison (per site)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'bess_phys_vs_virt.png', width_in=6.5,
        caption='Figure: Gross 6-month value capture for a physical 40 MW/160 MWh BESS vs an equivalent virtual TBx swap (x=4, daily), with the $3M/site lease (physical) or $3M/site fixed payment (virtual TBx) shown as the breakeven hurdle. At baseline drift, the gross capture is ~$1.0M (Houston) and ~$1.75M (West) — both well below the $3M cost — so neither contract pays at default rates.')
    # Orphan Valuation images
    anchor = insert_para_after(anchor, doc, text='Valuation diagnostics (additional)',
                                style='Heading 3')
    for key in ['valuation_1', 'valuation_2']:
        anchor = insert_image_para_after(anchor, doc, ORPHAN_IMAGES[key],
            width_in=6.0,
            caption='Figure: Valuation diagnostic (analysis text pending).')

    # ───────── Analyses > Power pricing ─────────
    anchor = targets[('Analyses', 'Power pricing')]
    # 1) LMP vs toll daily mean (image2) + prose
    anchor = insert_para_after(anchor, doc,
        text='Daily-mean LMP vs toll cost (last 3 months of 2025)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        WRITE_IMGS / 'image2.png', width_in=6.0,
        caption='Figure: Daily-mean LMP at HB_HOUSTON with p5-p95 hourly band, overlaid with implied toll cost under HSC (top) and Henry Hub (bottom) gas pricing conventions.')
    for para_text in WRITING_PROSE['lmp_vs_toll_daily']:
        anchor = insert_para_after(anchor, doc, text=para_text, style='Normal')
    # 2) Hourly LMP vs toll + win % (image8) + prose
    anchor = insert_para_after(anchor, doc,
        text='Hourly LMP vs toll cost and spark-spread win frequency',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        WRITE_IMGS / 'image8.png', width_in=6.0,
        caption='Figure: Left panel — mean hourly LMP vs toll cost (HSC). Right panel — percentage of hours where the toll beats the grid (positive spark spread), by hour of day.')
    for para_text in WRITING_PROSE['lmp_vs_toll_hourly']:
        anchor = insert_para_after(anchor, doc, text=para_text, style='Normal')
    # 3) Intermittency (image10) + prose
    anchor = insert_para_after(anchor, doc,
        text='Renewable intermittency: HB_HOUSTON upper-tail quantiles and HB_WEST negative-LMP frequency',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        WRITE_IMGS / 'image10.png', width_in=6.0,
        caption='Figure: Left — distribution of hourly LMPs at HB_HOUSTON by hour of day with quantile bands. Right — share of hours with negative LMP at HB_WEST by month.')
    for para_text in WRITING_PROSE['intermittency']:
        anchor = insert_para_after(anchor, doc, text=para_text, style='Normal')
    # 4) MC validation (image11) + prose
    anchor = insert_para_after(anchor, doc,
        text='Forward Monte-Carlo simulation paths (June–November 2026)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        WRITE_IMGS / 'image11.png', width_in=6.5,
        caption='Figure: 5,000-path Monte-Carlo simulations for four state variables — hourly LMP at HB_HOUSTON and HB_WEST and daily HH and HSC natural gas prices. 50 representative paths shown with 5th-95th percentile fan; dashed red line marks historical mean.')
    for para_text in WRITING_PROSE['mc_validation']:
        anchor = insert_para_after(anchor, doc, text=para_text, style='Normal')

    # ───────── Analyses > Tokens/compute ─────────
    anchor = targets[('Analyses', 'Tokens/compute')]
    # 1) Full cadence ranking table (Phase A)
    anchor = insert_para_after(anchor, doc,
        text='Full Phase A cadence ranking',
        style='Heading 3')
    cadence_rows = [
        ['Cadence', 'Mean profit ($M)', 'Std ($M)', 'p5 ($M)', 'p95 ($M)'],
        ['25d',  '15,645.66', '1.24', '15,643.62', '15,647.22'],
        ['30d',  '19,747.15', '1.24', '19,745.11', '19,748.72'],
        ['45d',  '27,067.02', '1.24', '27,064.98', '27,068.59'],
        ['60d',  '31,209.18', '1.24', '31,207.14', '31,210.74'],
        ['63d',  '31,873.46', '1.24', '31,871.42', '31,875.03'],
        ['74d',  '33,976.14', '1.24', '33,974.10', '33,977.71'],
        ['75d',  '34,142.46', '1.24', '34,140.42', '34,144.03'],
        ['85d',  '35,586.74', '1.24', '35,584.70', '35,588.31'],
        ['90d ★', '36,372.21', '1.24', '36,370.17', '36,373.78'],
    ]
    parent_proxy = anchor._parent
    cad_table = insert_table_after(anchor, doc, cadence_rows,
                                    widths_in=[0.9, 1.4, 0.8, 1.3, 1.3])
    anchor = append_after_table(cad_table, parent_proxy, doc,
        text='Table: Phase A mean profit across 9 candidate cadences (filtered to those satisfying the 500 MWh-grid/day RFP training floor). The 25d→90d sweep traces a near-monotonic profit gain of $20.7B (★ marks the Stage-2 verified winner). Standard deviations are tight (~$1.2M against ~$20-36B means) because the MC paths share calibration and the cadence-vs-cadence gap is dominated by the deterministic compute-allocation accounting, not the stochastic price processes.',
        style='Normal', italic=True, size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # 2) Verification: cadence winner under Phase C optimal procurement
    anchor = insert_para_after(anchor, doc,
        text='Verification: cadence winner under Phase C optimal procurement (LMP-only)',
        style='Heading 3')
    verif_rows = [
        ['Cadence', 'Mean profit ($M) under LMP-only', 'Δ vs 90d'],
        ['63d',  '31,878.51', '−4,498.75'],
        ['74d',  '33,981.19', '−2,396.07'],
        ['85d',  '35,591.79', '−785.47'],
        ['90d ★', '36,377.26', '0.00'],
    ]
    parent_proxy = anchor._parent
    verif_table = insert_table_after(anchor, doc, verif_rows,
                                      widths_in=[1.0, 2.6, 1.4])
    anchor = append_after_table(verif_table, parent_proxy, doc,
        text='Table: Re-solving the Phase A ±30% cadence neighborhood under the Phase C winning procurement (LMP-only) confirms the cadence winner does not shift when negative-NPV options are stripped. Cadence-vs-cadence profit gaps (~$0.8B–$4.5B) dwarf procurement-vs-procurement gaps (~$5M) by 3 orders of magnitude — the cadence decision is the dominant lever.',
        style='Normal', italic=True, size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # 3) Inference revenue rate over the horizon — shows the doc_blended × 0.25 dynamics
    anchor = insert_para_after(anchor, doc,
        text='Per-hour inference revenue rate over the 6-month horizon',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'inference_revenue_rate_horizon.png', width_in=6.5,
        caption='Figure: Hourly inference revenue rate ($K per grid-MWh of inference) over the Jun-Nov 2026 horizon under the doc_blended scheme with the 0.25 global multiplier. Red dotted lines mark release dates (R1 / R2 / R3 at 90d cadence). Each release resets the 270-day decay clock and applies the per-release uplift factor (1.346× per release at 90d); the global 0.25 multiplier is layered on top. The visible sawtooth amplitude (~5%) is small relative to the absolute rate because the 270-day halflife is long vs the 6-month horizon.')

    # ───────── Analyses > BESS/Nat-gas Tolling Optionality ─────────
    anchor = targets[('Analyses', 'BESS/Nat-gas Tolling Optionality')]
    # 1) Cross-Method Comparison (relocated from prior standalone slot) — toll valuation
    anchor = insert_para_after(anchor, doc,
        text='Cross-Method Comparison: Data-Center LP vs Power-Pricing-Only MC',
        style='Heading 3')
    for p_src in cross_method_paras:
        anchor = insert_para_after(anchor, doc, text=p_src.text, style='Normal')
    # Grouped bar chart visualizing the three valuation methods × two stress states
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'cross_method_comparison.png', width_in=6.5,
        caption='Figure: Gross toll value at 100 MW reservation over the 6-month horizon, three valuation methods × two stress states. M1 — external pricing-only MC (10K paths, MLE σ); M2 — our spark-spread MC (50 paths, tail_q σ); M3 — our LP joint optimization (50 paths, tail_q σ; cadence-invariant). All three methods agree on direction and order of magnitude; M2/M3 cluster within ~$0.1M of each other in the no-stress baseline and ~$0.25M apart under Uri, validating the LP\'s toll valuation against a price-only spark-spread benchmark.')
    # Compact summary table of the same numbers
    cm_rows = [
        ['Method', 'Paths × σ basis', 'No-stress toll value ($M)', 'Uri-stress toll value ($M)'],
        ['M1 — pricing-only MC (external)', '10K × MLE σ',  '1.42', '4.70'],
        ['M2 — our spark-spread MC',         '50 × tail_q σ', '2.27', '6.83'],
        ['M3 — our LP joint optimization',  '50 × tail_q σ', '2.18', '7.07'],
    ]
    parent_proxy = anchor._parent
    cm_table = insert_table_after(anchor, doc, cm_rows,
                                   widths_in=[2.4, 1.5, 1.5, 1.5])
    anchor = append_after_table(cm_table, parent_proxy, doc,
        text='Table: Three methods consistently estimate toll gross value at 100 MW × 6 months. M3 is our headline number used in the procurement decision (matches Δ vs LMP-only at K=$0 in the corrected fine-grid sweep). The M2 → M3 closeness is the key sanity check: a price-only spark-spread valuation lands within 5% of the LP\'s joint optimization, indicating the LP is not extracting non-physical option value.',
        style='Normal', italic=True, size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    # 2) IMHR distribution (image1) + prose
    anchor = insert_para_after(anchor, doc,
        text='Implied Market Heat Rate distribution at HB_HOUSTON',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        WRITE_IMGS / 'image1.png', width_in=6.0,
        caption='Figure: Distribution of implied market heat rate (IMHR) at HB_HOUSTON over the last three months of 2025. Left — HSC gas basis; right — Henry Hub basis. Reference lines: CCGT (6.8 MMBtu/MWh), Antelope peaker (9.2), simple-cycle peaker (9.5, our toll heat rate).')
    for para_text in WRITING_PROSE['imhr_distribution']:
        anchor = insert_para_after(anchor, doc, text=para_text, style='Normal')
    # NEW: Monthly 100-MW spark spread option values (RFP-explicit deliverable)
    anchor = insert_para_after(anchor, doc,
        text='Monthly 100-MW spark spread option value (Jun-Nov 2026)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'monthly_spark_spread_options.png', width_in=6.5,
        caption='Figure: 6-month monthly spark spread option value at HB_HOUSTON vs Henry Hub, 50 MC paths, 100-MW notional. Left panel — strike = 9.5 × HH (heat-rate-only, ignores VOM adder). Right panel — strike = 9.5 × (HH + $3) (toll-equivalent, includes the RFP $3/MMBtu VOM premium). Error bars are 5th-95th percentile across paths. The toll-equivalent strike isolates the value the LP captures by exercising the physical Houston toll option.')
    # 2) Capacity payment K sweep, two-panel
    anchor = insert_para_after(anchor, doc,
        text='Capacity-payment sensitivity (no-stress vs Uri-stress)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'capacity_payment_sweep_two_panel_fine.png', width_in=6.8,
        caption='Figure: Dual-axis K sweep at 90d cadence (fine MW grid: 0,10,…,100; K grid 0.02 step). Blue step (left axis) — optimal MW reservation, the buyer\'s rational best response. Green/red dashed (right axis) — net profit vs LMP-only for optimal-MW and fixed-100MW respectively. Left panel: no-stress baseline; toll wins +$2.18M at K=$0, breakeven K* ≈ $3.6/kW-mo, well below the seller\'s $8/kW-mo anchor, so optimal MW collapses to 0 before the seller rate. Right panel: under uri_full overlay, toll wins +$7.07M at K=$0 and K* climbs to ~$11.8/kW-mo, so the toll remains signable above the seller rate. The transition itself is a narrow stair-step (100→80→50→20→0 across a ~$0.06/kW-mo band) rather than a single bang-bang jump.')
    # 3) Cross-snapshot drift summary table — 16 rows (4 drifts × 4 stresses)
    anchor = insert_para_after(anchor, doc,
        text='Cross-snapshot drift × stress summary (16 scenarios)',
        style='Heading 3')
    table_rows = [
        ['Stress', 'Drift scenario', 'Gas / Power drift', 'Cadence', 'Procurement winner', 'Mean profit ($M)', 'Δ vs LMP-only ($M)'],
        ['none',     'baseline',      '0%  / 0%',     '90d', 'LMP only',    '36,377.26', '+0.00'],
        ['none',     'ai_structural', '0.5% / 1.0%',  '90d', 'LMP only',    '36,376.86', '+0.00'],
        ['none',     'mild_drift',    '3.0% / 1.5%',  '90d', 'LMP only',    '36,376.65', '+0.00'],
        ['none',     'ai_plus_brent', '6.5% / 4.0%',  '90d', 'LMP only',    '36,375.63', '+0.00'],
        ['mild',     'baseline',      '0%  / 0%',     '90d', 'LMP only',    '36,375.61', '+0.00'],
        ['mild',     'ai_structural', '0.5% / 1.0%',  '90d', 'LMP only',    '36,375.21', '+0.00'],
        ['mild',     'mild_drift',    '3.0% / 1.5%',  '90d', 'LMP only',    '36,375.00', '+0.00'],
        ['mild',     'ai_plus_brent', '6.5% / 4.0%',  '90d', 'LMP only',    '36,373.99', '+0.00'],
        ['moderate', 'baseline',      '0%  / 0%',     '90d', 'LMP only',    '36,374.06', '+0.00'],
        ['moderate', 'ai_structural', '0.5% / 1.0%',  '90d', 'LMP only',    '36,373.65', '+0.00'],
        ['moderate', 'mild_drift',    '3.0% / 1.5%',  '90d', 'LMP only',    '36,373.45', '+0.00'],
        ['moderate', 'ai_plus_brent', '6.5% / 4.0%',  '90d', 'LMP only',    '36,372.43', '+0.00'],
        ['uri_full', 'baseline',      '0%  / 0%',     '90d', 'LMP + toll', '36,365.67', '+2.27'],
        ['uri_full', 'ai_structural', '0.5% / 1.0%',  '90d', 'LMP + toll', '36,365.32', '+2.32'],
        ['uri_full', 'mild_drift',    '3.0% / 1.5%',  '90d', 'LMP + toll', '36,365.09', '+2.30'],
        ['uri_full', 'ai_plus_brent', '6.5% / 4.0%',  '90d', 'LMP + toll', '36,364.16', '+2.39'],
    ]
    parent_proxy = anchor._parent
    table = insert_table_after(anchor, doc, table_rows,
                                widths_in=[0.65, 1.25, 1.10, 0.55, 1.20, 1.05, 0.95])
    cap = append_after_table(table, parent_proxy, doc,
        text='Table: Final procurement policy across 4 drifts × 4 stress overlays at the LP-winner 90d cadence. The procurement winner is LMP-only for every (drift, stress) combination through mild and moderate stress; only under the Uri-stressed overlay does LMP+toll win, by $2.27–$2.39M (cadence-invariant — confirmed against 30d K-sweep CSVs). Drift-only sensitivity is small (≲$1.7M across drifts at any fixed stress). Stress overlay is the dominant axis: it shifts the headline mean profit by $1.6M / step (none→mild→moderate, mild gas tail) and by ≈$10M and flips the winner from none→uri_full.',
        style='Normal', italic=True, size_pt=10, alignment=WD_ALIGN_PARAGRAPH.CENTER)
    anchor = cap

    # 3b) Procurement delta heatmaps — full-cost + variable-cost view
    anchor = insert_para_after(anchor, doc,
        text='Procurement delta heatmaps (16 scenarios × 7 procurement choices)',
        style='Heading 3')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'procurement_delta_heatmap_full_cost.png', width_in=6.8,
        caption='Figure (full-cost view): Δ profit vs LMP-only for each procurement option across all 16 (drift × stress) snapshots, evaluated at the seller\'s default K=$8/kW-mo and 30d K-sweep cadence (deltas are cadence-invariant — confirmed identical to 4 decimal places against 90d). Green cells = option beats LMP-only; red = loses. Only the uri_full rows show any positive deltas, and only for toll-containing options. The toll-only column shows the classic stress ladder: −$2.6M (no-stress) → −$2.0M (mild) → −$1.3M (moderate) → +$2.3M (Uri). BESS-only options never beat LMP-only at the default lease.')
    anchor = insert_image_para_after(anchor, doc,
        ASSETS / 'procurement_delta_heatmap_variable_cost.png', width_in=6.8,
        caption='Figure (variable-cost view): same axes, but with the fixed capacity-payment ($4.8M) and BESS-lease ($3M/site) added back so cells reflect operational value only. Every cell is positive: under variable-cost framing the LP loves the dispatch flexibility of every procurement option. Toll+BESS-both is best in every row (peaks at +$11.6M under Uri / ai_plus_brent). The gap between this heatmap and the full-cost view is the lease tax the operator pays for that flexibility.')
    # 4) Orphan Risk Analysis image
    anchor = insert_para_after(anchor, doc, text='Risk-analysis diagnostic (additional)',
                                style='Heading 3')
    anchor = insert_image_para_after(anchor, doc, ORPHAN_IMAGES['risk_analysis'],
        width_in=6.0,
        caption='Figure: Risk-analysis diagnostic (analysis text pending).')

    # 3) Remove the old Cross-Method block (heading + body paragraphs we relocated)
    print('Removing relocated Cross-Method Comparison block from original slot…')
    old_cross = targets['cross_method_heading']
    old_cross._element.getparent().remove(old_cross._element)
    for p in cross_method_paras:
        if p._element.getparent() is not None:
            p._element.getparent().remove(p._element)

    # 4) Save
    print(f'Saving -> {DST_REPORT}')
    doc.save(str(DST_REPORT))
    print('Done.')


if __name__ == '__main__':
    main()
